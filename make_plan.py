"""Generate the project plan as a .docx using python-docx (same style as the report)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = r"d:\Рабочий стол\Homework\llm final project\2024998004014_Anton Volkov"
os.makedirs(OUT_DIR, exist_ok=True)
OUT = os.path.join(OUT_DIR, "Project_Plan.docx")

INK = RGBColor(0x2E, 0x2A, 0x26)
SAGE = RGBColor(0x3F, 0x5A, 0x4A)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = INK
for h, size in (("Heading 1", 15), ("Heading 2", 12)):
    st = doc.styles[h]
    st.font.name = "Calibri"
    st.font.color.rgb = SAGE
    st.font.size = Pt(size)


def body(text, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(after)
    return p


def day(label, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label + " — ")
    r.bold = True
    p.add_run(text)
    return p


def milestone(text):
    p = doc.add_paragraph()
    r = p.add_run("Milestone: " + text)
    r.italic = True
    r.font.color.rgb = SAGE
    p.paragraph_format.space_after = Pt(10)
    return p


# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Project Plan — AI Agent Vocabulary Assistant (Lexa)")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = INK

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Author: Anton Volkov   |   Student ID: 2024998004014   |   Duration: 3 weeks")

body("A vocabulary learning system built with LLMs, an AI agent workflow, OpenClaw, "
     "AI coding tools and cloud deployment. The plan is organized into three one-week "
     "phases, broken down day by day.", italic=True, after=12)

# Week 1
doc.add_heading("Week 1 — Foundation & Core Agents", level=1)
body("Goal: a working backend with the two AI agents.", italic=True, after=4)
day("Day 1", "Requirements analysis; fix the tech stack (Qwen LLM, Tavily search, OpenClaw, "
    "Node/Express/Prisma/PostgreSQL, Next.js, Railway/Vercel); create repository and project structure.")
day("Day 2", "Backend skeleton: Express + TypeScript, Prisma + PostgreSQL (Docker), database schema "
    "(User / Word / Example), first migration, /health endpoint.")
day("Day 3", "Services: Qwen LLM client (JSON output), Tavily news-search client (restricted to news "
    "domains), Zod validation schemas.")
day("Day 4", "Example Search Agent: Tavily search -> LLM selects the best sentence -> LLM translates to "
    "Chinese -> persist word + example + source + URL.")
day("Day 5", "Vocabulary Tutor Agent: phonetic (IPA), part of speech, Chinese meaning, collocations, "
    "synonyms, antonyms.")
day("Day 6-7", "REST API (/api/words: add / list / get / delete / review); end-to-end testing; buffer.")
milestone("a word added via the API returns a real news example, a translation and a full dictionary entry.")

# Week 2
doc.add_heading("Week 2 — Surfaces (Bot + Web) & OpenClaw", level=1)
body("Goal: both user surfaces working on one shared backend.", italic=True, after=4)
day("Day 8", "Telegram bot baseline (add / list commands).")
day("Day 9", "OpenClaw integration: install, custom vocab-assistant skill, connect Telegram, route "
    "commands to the REST API (per-user by Telegram ID).")
day("Day 10", "Frontend scaffold: Next.js 15 + Tailwind + TanStack Query; typed API client.")
day("Day 11", "Web pages: word list, word detail, flashcard review.")
day("Day 12", "Features: search, delete, due-for-review badges, spaced repetition; multi-user accounts.")
day("Day 13", "UI redesign (editorial “Lexa” design system) + Recall-check quiz + animations.")
day("Day 14", "Integration testing bot <-> web <-> backend; bug fixes.")
milestone("a word added in Telegram immediately appears on the web app.")

# Week 3
doc.add_heading("Week 3 — Deployment, Polish & Deliverables", level=1)
body("Goal: a publicly deployed system and the full submission package.", italic=True, after=4)
day("Day 15", "Deployment prep: environment config, build/migrate scripts, .gitignore, push to GitHub.")
day("Day 16", "Deploy backend + PostgreSQL (Railway) and frontend (Vercel); wire public URLs.")
day("Day 17", "Containerize and deploy OpenClaw (Docker on Railway, always-on); configure agent auth.")
day("Day 18", "Vision feature (add a word from a photo, Qwen3-VL); seed demo content; desktop + mobile testing.")
day("Day 19", "Project document: overview, tech stack, agent workflow + diagram, screenshots, AI-usage summary.")
day("Day 20", "Record the 3-5 minute demo video.")
day("Day 21", "Final review against the grading rubric; package deliverables (document + video + source); submit.")
milestone("system publicly accessible on desktop and mobile; all deliverables submitted.")

doc.save(OUT)
print("WROTE", OUT)
